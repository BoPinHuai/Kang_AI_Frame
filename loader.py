import csv
import sys
from pathlib import Path
import pypdf
from docx import Document

# Ansys 等大型 PDF 有极深的对象引用链，需要提高递归上限
sys.setrecursionlimit(10000)

SUPPORTED = {".pdf", ".docx", ".md", ".xlsx", ".csv"}


# ── 对外主接口：返回带位置标签的分段 ──────────────────────────
def load_segments(file_path: str) -> list[dict]:
    """统一读取文档，返回 [{"text": str, "loc": str}, ...]

    loc 是来源位置标签（前端显示）：
      PDF   → "第3页"      （按页分段）
      Excel → "工作表 Sheet1"（按工作表分段）
      docx 表格 → "表格"
      md / docx 正文 / csv → ""（无位置概念）
    """
    suffix = Path(file_path).suffix.lower()
    if suffix == ".pdf":
        return _seg_pdf(file_path)
    elif suffix == ".docx":
        return _seg_docx(file_path)
    elif suffix == ".md":
        return [{"text": _read_text(file_path), "loc": ""}]
    elif suffix == ".xlsx":
        return _seg_xlsx(file_path)
    elif suffix == ".csv":
        return _seg_csv(file_path)
    else:
        raise ValueError(f"不支持的格式: {suffix}（支持 {', '.join(sorted(SUPPORTED))}）")


def load_document(file_path: str) -> str:
    """兼容旧接口：把所有分段拼成一个字符串。"""
    return "\n\n".join(s["text"] for s in load_segments(file_path) if s["text"].strip())


# ── 各格式分段实现 ───────────────────────────────────────────
def _seg_pdf(file_path: str) -> list[dict]:
    reader = pypdf.PdfReader(file_path)
    segs = []
    for i, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        if text.strip():
            segs.append({"text": text, "loc": f"第{i}页"})
    return segs


def _seg_docx(file_path: str) -> list[dict]:
    doc = Document(file_path)
    segs = []
    # 正文段落
    body = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if body.strip():
        segs.append({"text": body, "loc": ""})
    # 表格（python-docx 可读取 .docx 内嵌表格）
    for tbl in doc.tables:
        rows = [[c.text.strip() for c in row.cells] for row in tbl.rows]
        rows = [r for r in rows if any(r)]
        if len(rows) >= 1:
            text = _rows_to_text(rows[0], rows[1:]) if len(rows) > 1 else "，".join(rows[0])
            if text.strip():
                segs.append({"text": text, "loc": "表格"})
    return segs or [{"text": "", "loc": ""}]


def _seg_xlsx(file_path: str) -> list[dict]:
    from openpyxl import load_workbook
    wb = load_workbook(file_path, read_only=True, data_only=True)
    segs = []
    for ws in wb.worksheets:
        rows = []
        for r in ws.iter_rows(values_only=True):
            cells = ["" if c is None else str(c).strip() for c in r]
            if any(cells):
                rows.append(cells)
        if not rows:
            continue
        text = _rows_to_text(rows[0], rows[1:]) if len(rows) > 1 else "，".join(rows[0])
        if text.strip():
            segs.append({"text": text, "loc": f"工作表 {ws.title}"})
    wb.close()
    return segs


def _seg_csv(file_path: str) -> list[dict]:
    rows = []
    with open(file_path, newline="", encoding="utf-8-sig", errors="ignore") as f:
        for r in csv.reader(f):
            cells = [c.strip() for c in r]
            if any(cells):
                rows.append(cells)
    if not rows:
        return [{"text": "", "loc": ""}]
    text = _rows_to_text(rows[0], rows[1:]) if len(rows) > 1 else "，".join(rows[0])
    return [{"text": text, "loc": ""}]


# ── 表格 → 行式文本（每行自带表头，避免切块后丢表头）──────────
def _rows_to_text(headers: list[str], data_rows: list[list[str]]) -> str:
    """把表格转成「表头：值」的行式文本，每行之间空行分隔，
    便于分块时保留完整行、且每行可独立检索。"""
    lines = []
    for row in data_rows:
        pairs = []
        for h, v in zip(headers, row):
            v = (v or "").strip()
            if v:
                pairs.append(f"{h}：{v}" if h else v)
        if pairs:
            lines.append("，".join(pairs))
    return "\n\n".join(lines)


def _read_text(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


# ── 旧批量接口（保留兼容）──────────────────────────────────
def load_all_docs(folder: str) -> list[dict]:
    docs = []
    root = Path(folder)
    for file in root.rglob("*"):
        if not file.is_file():
            continue
        if file.name.startswith("~$") or file.name.startswith("."):
            continue
        if file.suffix.lower() in SUPPORTED:
            rel = str(file.relative_to(root)).replace("\\", "/")
            print(f"  正在读取: {rel}")
            text = load_document(str(file))
            if text.strip():
                docs.append({"filename": rel, "text": text})
            else:
                print(f"  警告: {rel} 内容为空，已跳过")
    return docs
